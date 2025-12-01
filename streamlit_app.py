import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import io
from PIL import Image, ImageOps, ImageDraw
import pypdf
from datetime import datetime
import json
import urllib.parse
# Gestione import SerpApi per evitare crash se manca la lib nel virtual env locale (su cloud ci sarà)
try:
    from serpapi import GoogleSearch
except ImportError:
    GoogleSearch = None

# -----------------------------------------------------------------------------
# 1. CONFIGURAZIONE
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Global Career Coach", layout="wide", initial_sidebar_state="expanded")

# -----------------------------------------------------------------------------
# 2. CSS
# -----------------------------------------------------------------------------
st.markdown("""
    <style>
    div[data-baseweb="select"] > div { cursor: pointer !important; }
    button { cursor: pointer !important; }
    .job-card {
        padding: 15px;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        background-color: #ffffff;
        margin-bottom: 12px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .job-title {
        font-size: 18px;
        font-weight: bold;
        color: #20547D;
    }
    .job-company {
        font-size: 14px;
        color: #555;
        margin-bottom: 8px;
    }
    </style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 3. STATE MANAGEMENT
# -----------------------------------------------------------------------------
if 'lang_code' not in st.session_state: st.session_state['lang_code'] = 'it'
if 'generated_data' not in st.session_state: st.session_state['generated_data'] = None
if 'processed_photo' not in st.session_state: st.session_state['processed_photo'] = None
if 'job_search_results' not in st.session_state: st.session_state['job_search_results'] = None

# -----------------------------------------------------------------------------
# 4. COSTANTI E TRADUZIONI
# -----------------------------------------------------------------------------
LANG_DISPLAY = {
    "Italiano": "it", "English (US)": "en_us", "English (UK)": "en_uk",
    "Deutsch (Deutschland)": "de_de", "Deutsch (Schweiz)": "de_ch",
    "Français": "fr", "Español": "es", "Português": "pt"
}

TRANSLATIONS = {
    'it': {'sidebar_title': 'Impostazioni Profilo', 'lang_label': 'Lingua', 'photo_label': 'Foto Profilo', 'border_label': 'Bordo (px)', 'preview_label': 'Anteprima', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Carica CV (PDF)', 'upload_help': 'Trascina file qui', 'step2_title': '2. Annuncio di Lavoro', 'job_placeholder': 'Incolla qui il testo dell\'offerta...', 'btn_label': 'Genera Documenti', 'spinner_msg': 'Elaborazione in corso...', 'tab_cv': 'CV Generato', 'tab_letter': 'Lettera', 'down_cv': 'Scarica CV (Word)', 'down_let': 'Scarica Lettera (Word)', 'success': 'Fatto!', 'error': 'Errore', 'profile_title': 'PROFILO PERSONALE', 'search_role': 'Che lavoro cerchi?', 'search_loc': 'Dove?', 'search_rad': 'Raggio (km)', 'search_btn': 'Trova Lavori 🔎', 'search_res_title': 'Offerte Trovate:', 'search_info': 'Ecco le offerte reali trovate su Google Jobs.', 'no_jobs': 'Nessun lavoro trovato.', 'upload_first': '⚠️ Carica prima il CV!', 'tab_search': '🌍 Ricerca Lavoro', 'tab_docs': '📄 Documenti', 'apply': 'Candidati Ora 🚀', 'source': 'Fonte'},
    'en_us': {'sidebar_title': 'Profile Settings', 'lang_label': 'Language', 'photo_label': 'Profile Photo', 'border_label': 'Border (px)', 'preview_label': 'Preview', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Upload CV (PDF)', 'upload_help': 'Drop file here', 'step2_title': '2. Job Advertisement', 'job_placeholder': 'Paste job offer...', 'btn_label': 'Generate Documents', 'spinner_msg': 'Processing...', 'tab_cv': 'Generated CV', 'tab_letter': 'Cover Letter', 'down_cv': 'Download CV', 'down_let': 'Download Letter', 'success': 'Done!', 'error': 'Error', 'profile_title': 'PROFESSIONAL PROFILE', 'search_role': 'Job Title', 'search_loc': 'Location', 'search_rad': 'Radius (km)', 'search_btn': 'Find Jobs 🔎', 'search_res_title': 'Found Jobs:', 'search_info': 'Real jobs found on Google Jobs.', 'no_jobs': 'No jobs found.', 'upload_first': '⚠️ Upload CV first!', 'tab_search': '🌍 Job Search', 'tab_docs': '📄 Documents', 'apply': 'Apply Now 🚀', 'source': 'Source'},
    'de_ch': {'sidebar_title': 'Einstellungen', 'lang_label': 'Sprache', 'photo_label': 'Profilbild', 'border_label': 'Rahmen (px)', 'preview_label': 'Vorschau', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Lebenslauf hochladen (PDF)', 'upload_help': 'Datei hier ablegen', 'step2_title': '2. Stelleninserat', 'job_placeholder': 'Stelleninserat hier einfügen...', 'btn_label': 'Dokumente erstellen', 'spinner_msg': 'Verarbeitung läuft...', 'tab_cv': 'Lebenslauf', 'tab_letter': 'Motivationsschreiben', 'down_cv': 'Lebenslauf laden', 'down_let': 'Brief laden', 'success': 'Fertig!', 'error': 'Fehler', 'profile_title': 'PERSÖNLICHES PROFIL', 'search_role': 'Welcher Job?', 'search_loc': 'Wo?', 'search_rad': 'Umkreis (km)', 'search_btn': 'Jobs suchen 🔎', 'search_res_title': 'Gefundene Jobs:', 'search_info': 'Echte Jobs gefunden auf Google Jobs.', 'no_jobs': 'Keine Jobs gefunden.', 'upload_first': '⚠️ Zuerst Lebenslauf hochladen!', 'tab_search': '🌍 Jobsuche', 'tab_docs': '📄 Dokumente', 'apply': 'Jetzt Bewerben 🚀', 'source': 'Quelle'},
    'de_de': {'sidebar_title': 'Einstellungen', 'lang_label': 'Sprache', 'photo_label': 'Profilbild', 'border_label': 'Rahmen (px)', 'preview_label': 'Vorschau', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Lebenslauf hochladen (PDF)', 'upload_help': 'Datei hier ablegen', 'step2_title': '2. Stellenanzeige', 'job_placeholder': 'Stellenanzeige einfügen...', 'btn_label': 'Dokumente erstellen', 'spinner_msg': 'Verarbeitung läuft...', 'tab_cv': 'Lebenslauf', 'tab_letter': 'Anschreiben', 'down_cv': 'Lebenslauf laden', 'down_let': 'Brief laden', 'success': 'Fertig!', 'error': 'Fehler', 'profile_title': 'PERSÖNLICHES PROFIL', 'search_role': 'Welcher Job?', 'search_loc': 'Wo?', 'search_rad': 'Umkreis (km)', 'search_btn': 'Jobs suchen 🔎', 'search_res_title': 'Gefundene Jobs:', 'search_info': 'Echte Jobs gefunden auf Google Jobs.', 'no_jobs': 'Keine Jobs gefunden.', 'upload_first': '⚠️ Zuerst Lebenslauf hochladen!', 'tab_search': '🌍 Jobsuche', 'tab_docs': '📄 Dokumente', 'apply': 'Jetzt Bewerben 🚀', 'source': 'Quelle'},
    'fr': {'sidebar_title': 'Paramètres', 'lang_label': 'Langue', 'photo_label': 'Photo', 'border_label': 'Bordure (px)', 'preview_label': 'Aperçu', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Télécharger CV', 'upload_help': 'Déposez ici', 'step2_title': '2. Offre d\'Emploi', 'job_placeholder': 'Collez l\'offre...', 'btn_label': 'Générer', 'spinner_msg': 'Traitement...', 'tab_cv': 'CV Généré', 'tab_letter': 'Lettre', 'down_cv': 'Télécharger CV', 'down_let': 'Télécharger Lettre', 'success': 'Terminé!', 'error': 'Erreur', 'profile_title': 'PROFIL', 'search_role': 'Quel emploi?', 'search_loc': 'Où?', 'search_rad': 'Rayon (km)', 'search_btn': 'Chercher 🔎', 'search_res_title': 'Résultats:', 'search_info': 'Offres réelles Google Jobs.', 'no_jobs': 'Aucun résultat.', 'upload_first': '⚠️ CV requis!', 'tab_search': '🌍 Recherche', 'tab_docs': '📄 Documents', 'apply': 'Postuler 🚀', 'source': 'Source'},
    'es': {'sidebar_title': 'Configuración', 'lang_label': 'Idioma', 'photo_label': 'Foto', 'border_label': 'Borde (px)', 'preview_label': 'Vista', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Subir CV', 'upload_help': 'Arrastra aquí', 'step2_title': '2. Oferta de Empleo', 'job_placeholder': 'Pega oferta...', 'btn_label': 'Generar', 'spinner_msg': 'Procesando...', 'tab_cv': 'CV Generado', 'tab_letter': 'Carta', 'down_cv': 'Descargar CV', 'down_let': 'Descargar Carta', 'success': 'Hecho', 'error': 'Error', 'profile_title': 'PERFIL', 'search_role': '¿Puesto?', 'search_loc': '¿Dónde?', 'search_rad': 'Radio (km)', 'search_btn': 'Buscar 🔎', 'search_res_title': 'Resultados:', 'search_info': 'Ofertas reales Google Jobs.', 'no_jobs': 'Sin resultados.', 'upload_first': '⚠️ CV requerido!', 'tab_search': '🌍 Buscar Empleo', 'tab_docs': '📄 Documentos', 'apply': 'Aplicar 🚀', 'source': 'Fuente'},
    'pt': {'sidebar_title': 'Configurações', 'lang_label': 'Idioma', 'photo_label': 'Foto', 'border_label': 'Borda (px)', 'preview_label': 'Visualizar', 'main_title': 'Global Career Coach 🌍', 'step1_title': '1. Carregar CV', 'upload_help': 'Arraste aqui', 'step2_title': '2. Anúncio', 'job_placeholder': 'Cole anúncio...', 'btn_label': 'Gerar', 'spinner_msg': 'Processando...', 'tab_cv': 'CV Gerado', 'tab_letter': 'Carta', 'down_cv': 'Baixar CV', 'down_let': 'Baixar Carta', 'success': 'Pronto', 'error': 'Erro', 'profile_title': 'PERFIL', 'search_role': 'Cargo?', 'search_loc': 'Onde?', 'search_rad': 'Raio (km)', 'search_btn': 'Buscar 🔎', 'search_res_title': 'Resultados:', 'search_info': 'Vagas reais Google Jobs.', 'no_jobs': 'Sem resultados.', 'upload_first': '⚠️ CV necessário!', 'tab_search': '🌍 Vagas', 'tab_docs': '📄 Documentos', 'apply': 'Candidatar 🚀', 'source': 'Fonte'}
}

SECTION_TITLES = {
    'it': {'experience': 'ESPERIENZA PROFESSIONALE', 'education': 'ISTRUZIONE', 'skills': 'COMPETENZE', 'languages': 'LINGUE', 'interests': 'INTERESSI', 'personal_info': 'DATI PERSONALI', 'profile_summary': 'PROFILO PERSONALE'},
    'de_ch': {'experience': 'BERUFSERFAHRUNG', 'education': 'AUSBILDUNG', 'skills': 'KENNTNISSE', 'languages': 'SPRACHEN', 'interests': 'INTERESSEN', 'personal_info': 'PERSÖNLICHE DATEN', 'profile_summary': 'PERSÖNLICHES PROFIL'},
    'de_de': {'experience': 'BERUFSERFAHRUNG', 'education': 'AUSBILDUNG', 'skills': 'KENNTNISSE', 'languages': 'SPRACHEN', 'interests': 'INTERESSEN', 'personal_info': 'PERSÖNLICHE DATEN', 'profile_summary': 'PERSÖNLICHES PROFIL'},
    'fr': {'experience': 'EXPÉRIENCE PROFESSIONNELLE', 'education': 'FORMATION', 'skills': 'COMPÉTENCES', 'languages': 'LANGUES', 'interests': 'INTÉRÊTS', 'personal_info': 'INFORMATIONS PERSONNELLES', 'profile_summary': 'PROFIL PROFESSIONNEL'},
    'en_us': {'experience': 'PROFESSIONAL EXPERIENCE', 'education': 'EDUCATION', 'skills': 'SKILLS', 'languages': 'LANGUAGES', 'interests': 'INTERESTS', 'personal_info': 'PERSONAL DETAILS', 'profile_summary': 'PROFESSIONAL PROFILE'},
    'en_uk': {'experience': 'WORK EXPERIENCE', 'education': 'EDUCATION', 'skills': 'SKILLS', 'languages': 'LANGUAGES', 'interests': 'INTERESTS', 'personal_info': 'PERSONAL DETAILS', 'profile_summary': 'PROFILE'},
    'es': {'experience': 'EXPERIENCIA LABORAL', 'education': 'EDUCACIÓN', 'skills': 'HABILIDADES', 'languages': 'IDIOMAS', 'interests': 'INTERESES', 'personal_info': 'DATOS PERSONALES', 'profile_summary': 'PERFIL PROFESIONAL'},
    'pt': {'experience': 'EXPERIÊNCIA PROFISSIONAL', 'education': 'EDUCAÇÃO', 'skills': 'COMPETÊNCIAS', 'languages': 'IDIOMAS', 'interests': 'INTERESSES', 'personal_info': 'DADOS PESSOAIS', 'profile_summary': 'PERFIL PROFISSIONAL'}
}

# -----------------------------------------------------------------------------
# 5. FUNZIONI HELPER
# -----------------------------------------------------------------------------

def get_todays_date(lang_code):
    now = datetime.now()
    if lang_code in ['de_ch', 'de_de', 'it', 'fr', 'es', 'pt']:
        return now.strftime("%d.%m.%Y")
    return now.strftime("%B %d, %Y")

def extract_text_from_pdf(uploaded_file):
    try:
        reader = pypdf.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except:
        return ""

def process_image(uploaded_img, border_size):
    try:
        img = Image.open(uploaded_img).convert("RGBA")
        size = (min(img.size), min(img.size))
        mask = Image.new('L', size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0) + size, fill=255)
        output = ImageOps.fit(img, size, centering=(0.5, 0.5))
        output.putalpha(mask)
        if border_size > 0:
            final_size = (size[0] + border_size * 2, size[1] + border_size * 2)
            bg = Image.new('RGBA', final_size, (0,0,0,0))
            draw_bg = ImageDraw.Draw(bg)
            draw_bg.ellipse((0, 0) + final_size, fill=(255,255,255,255))
            bg.paste(output, (border_size, border_size), output)
            return bg
        else:
            return output
    except:
        return None

def set_table_background(table, color_hex):
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="20547D"/></w:pBdr>'.format(nsdecls('w')))
    pPr.append(pBdr)

# -----------------------------------------------------------------------------
# 6. RICERCA LAVORO (SERPAPI PURO - NO AI)
# -----------------------------------------------------------------------------
def search_jobs_master(role, location, radius, lang_ui):
    if "SERPAPI_API_KEY" not in st.secrets:
        st.error("SERPAPI_API_KEY missing.")
        return []
    
    # Mappatura semplice per il parametro 'gl' (Paese) in base alla lingua,
    # migliorabile con un input "Nazione" separato, ma qui usiamo default CH per semplicità
    # o intuiamo dalla lingua.
    gl_code = "ch" 
    if "de_de" in lang_ui: gl_code = "de"
    elif "en_us" in lang_ui: gl_code = "us"
    elif "en_uk" in lang_ui: gl_code = "uk"
    elif "it" in lang_ui: gl_code = "it" # O 'ch' se si cerca in Ticino

    # Override: Se l'utente scrive "Zurich, CH", SerpApi capisce la location.
    # Il parametro 'gl' aiuta per i risultati locali.
    
    params = {
        "engine": "google_jobs",
        "q": f"{role} {location}",
        "hl": lang_ui.split('_')[0], # 'it', 'en', 'de'
        "gl": gl_code, 
        "radius": radius,
        "api_key": st.secrets["SERPAPI_API_KEY"]
    }

    try:
        search = GoogleSearch(params)
        results = search.get_dict().get("jobs_results", [])
        final_res = []
        
        for job in results[:15]: # Prendiamo i primi 15
            
            # --- LOGICA LINK INFALLIBILE ---
            link = None
            
            # 1. Cerca nei bottoni "Apply" (Candidati)
            apply_options = job.get("apply_options", [])
            for option in apply_options:
                if option.get("link"):
                    link = option.get("link")
                    break # Trovato il primo link utile
            
            # 2. Se non c'è link diretto, usa il JOB ID per il link Google Jobs
            # Questo apre la scheda specifica dell'offerta su Google
            if not link and "job_id" in job:
                # Esempio URL: https://www.google.com/search?ibp=htl;jobs#fpstate=tldetail&htivrt=jobs&htidocid={JOB_ID}
                j_id = job["job_id"]
                link = f"https://www.google.com/search?ibp=htl;jobs#fpstate=tldetail&htivrt=jobs&htidocid={j_id}"
            
            # 3. Ultima spiaggia: share_link
            if not link:
                link = job.get("share_link")

            # Se abbiamo un link, aggiungiamo il risultato
            if link:
                final_res.append({
                    "company": job.get("company_name", "Azienda"),
                    "role_title": job.get("title", role),
                    "location": job.get("location", location),
                    "link": link,
                    "source": apply_options[0].get("title") if apply_options else "Google Jobs"
                })
        
        return final_res

    except Exception as e:
        st.error(f"Search Error: {e}")
        return []

# -----------------------------------------------------------------------------
# 7. GENERAZIONE CV (CONGELATO)
# -----------------------------------------------------------------------------
def create_cv_docx(data, photo_img, lang_code):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.5)
    
    set_table_background(table, "20547D")

    cell_img = table.cell(0, 0)
    cell_img.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if photo_img:
        p = cell_img.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        img_byte = io.BytesIO()
        photo_img.save(img_byte, format="PNG")
        run.add_picture(img_byte, width=Cm(3.5))
    
    cell_txt = table.cell(0, 1)
    cell_txt.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p1 = cell_txt.paragraphs[0]
    run1 = p1.add_run(f"{data.get('nome', '')}\n")
    run1.font.size = Pt(24)
    run1.font.color.rgb = RGBColor(255, 255, 255)
    run1.bold = True
    
    info_text = f"{data.get('indirizzo','')}\n{data.get('telefono','')} | {data.get('email','')}"
    p2 = cell_txt.add_paragraph(info_text)
    run2 = p2.runs[0]
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(230, 230, 230)

    doc.add_paragraph().space_after = Pt(12)

    titles = SECTION_TITLES.get(lang_code, SECTION_TITLES['en_us'])
    
    if 'profile_summary' in data['cv_sections']:
        h = doc.add_paragraph(titles['profile_summary'])
        add_bottom_border(h)
        h.runs[0].bold = True
        h.runs[0].font.color.rgb = RGBColor(32, 84, 125)
        doc.add_paragraph(data['cv_sections']['profile_summary'].replace('**', ''))
        doc.add_paragraph("")

    sections = ['experience', 'education', 'skills', 'languages', 'interests']
    for key in sections:
        if key in data['cv_sections'] and data['cv_sections'][key]:
            h = doc.add_paragraph(titles[key])
            add_bottom_border(h)
            h.runs[0].bold = True
            h.runs[0].font.color.rgb = RGBColor(32, 84, 125)
            
            items = data['cv_sections'][key]
            if isinstance(items, list):
                for item in items:
                    p = doc.add_paragraph(item.replace('**', ''), style='List Bullet')
                    if key in ['experience', 'education']:
                        doc.add_paragraph("")
            else:
                doc.add_paragraph(str(items).replace('**', ''))
                doc.add_paragraph("")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_letter_docx(letter_data, personal_info, lang_code):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_paragraph(f"{personal_info.get('name')}\n{personal_info.get('address')}\n{personal_info.get('phone')}\n{personal_info.get('email')}")
    doc.add_paragraph("")
    doc.add_paragraph(get_todays_date(LANG_DISPLAY.get(lang_code, 'en')))
    doc.add_paragraph("")
    doc.add_paragraph(letter_data.get('recipient_block', ''))
    doc.add_paragraph("")
    p_obj = doc.add_paragraph(letter_data.get('subject_line', ''))
    p_obj.runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph(letter_data.get('body_content', ''))
    doc.add_paragraph("")
    doc.add_paragraph(letter_data.get('closing', 'Cordiali saluti'))
    doc.add_paragraph("\n\n")
    doc.add_paragraph(personal_info.get('name', ''))
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

# -----------------------------------------------------------------------------
# 8. AI GENERATION
# -----------------------------------------------------------------------------
def get_gemini_response(pdf_text, job_text, lang_code):
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel("models/gemini-3-pro-preview")
        prompt = f"""
        Act as an Expert Career Coach.
        Output Language: {lang_code} (Strictly).
        
        INPUT CV: {pdf_text[:4000]}
        INPUT JOB AD: {job_text}
        
        TASK:
        1. Extract personal info from CV.
        2. Create a Professional CV content adapted to the Job Ad.
        3. Create a Cover Letter adapted to the Job Ad.
        
        OUTPUT JSON FORMAT:
        {{
            "personal_info": {{"name": "...", "address": "...", "phone": "...", "email": "..."}},
            "cv_sections": {{
                "profile_summary": "...",
                "experience": ["Title - Company (Dates): Description", ...],
                "education": ["Title - Institute (Dates)", ...],
                "skills": ["Skill1", "Skill2", ...],
                "languages": ["Lang - Level", ...],
                "interests": "..."
            }},
            "letter_data": {{
                "recipient_block": "Hiring Manager...",
                "subject_line": "Application for...",
                "body_content": "Dear...",
                "closing": "Best regards"
            }}
        }}
        """
        response = model.generate_content(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception as e:
        return None

# -----------------------------------------------------------------------------
# 9. MAIN APP LOOP
# -----------------------------------------------------------------------------
def main():
    if 'lang_code' not in st.session_state: st.session_state.lang_code = "it"
    t_code = st.session_state.lang_code
    t = TRANSLATIONS[t_code]

    # Sidebar
    with st.sidebar:
        st.title(t['sidebar_title'])
        
        # Lingua
        lang_opts = list(LANG_DISPLAY.keys())
        # Trova indice corrente
        curr_idx = 0
        for k, v in LANG_DISPLAY.items():
            if v == t_code:
                curr_idx = lang_opts.index(k)
                break
        
        sel_lang = st.selectbox(t['lang_label'], lang_opts, index=curr_idx)
        st.session_state.lang_code = LANG_DISPLAY[sel_lang]
        
        # Foto
        st.divider()
        st.subheader(t['photo_label'])
        up_img = st.file_uploader("Upload", type=['jpg','png','jpeg'], label_visibility="collapsed")
        border = st.slider(t['border_label'], 0, 30, 5)
        
        if up_img:
            proc = process_image(up_img, border)
            st.session_state.processed_photo = proc
            st.image(proc, caption=t['preview_label'], width=150)

    # Main
    st.title(t['main_title'])
    
    # Tabs
    tab1, tab2 = st.tabs([t['tab_docs'], t['tab_search']])

    # TAB 1: Documenti
    with tab1:
        c1, c2 = st.columns(2)
        with c1:
            st.subheader(t['step1_title'])
            f_pdf = st.file_uploader("PDF", type=['pdf'], label_visibility="collapsed", key="cv_upl")
        with c2:
            st.subheader(t['step2_title'])
            job_txt = st.text_area("Job", placeholder=t['job_placeholder'], height=150, label_visibility="collapsed")

        if st.button(t['btn_label'], type="primary"):
            if f_pdf and job_txt:
                with st.spinner(t['spinner_msg']):
                    txt = extract_text_from_pdf(f_pdf)
                    data = get_gemini_response(txt, job_txt, st.session_state.lang_code)
                    if data:
                        st.session_state.generated_data = data
                        st.success(t['success'])
                    else:
                        st.error(t['error'])
            else:
                st.warning(t['upload_first'])
        
        if st.session_state.generated_data:
            st.divider()
            cd1, cd2 = st.columns(2)
            
            # CV
            docx_cv = create_cv_docx(st.session_state.generated_data, st.session_state.processed_photo, st.session_state.lang_code)
            with cd1:
                st.download_button(t['down_cv'], docx_cv.getvalue(), "CV.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            # Lettera
            docx_let = create_letter_docx(st.session_state.generated_data['letter_data'], st.session_state.generated_data['personal_info'], st.session_state.lang_code)
            with cd2:
                st.download_button(t['down_let'], docx_let.getvalue(), "Letter.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # TAB 2: Ricerca Lavoro (SOLO SERPAPI)
    with tab2:
        c_role, c_loc, c_rad = st.columns([2, 2, 1])
        role = c_role.text_input(t['search_role'])
        loc = c_loc.text_input(t['search_loc'], value="Zürich, CH")
        rad = c_rad.number_input(t['search_rad'], value=20)
        
        if st.button(t['search_btn']):
            with st.spinner(t['spinner_msg']):
                res = search_jobs_master(role, loc, rad, st.session_state.lang_code)
                st.session_state.job_search_results = res
        
        if st.session_state.job_search_results:
            st.success(f"{t['search_res_title']} {len(st.session_state.job_search_results)}")
            st.info(t['search_info'])
            
            for job in st.session_state.job_search_results:
                st.markdown(f"""
                <div class="job-card">
                    <div class="job-title">{job['role_title']}</div>
                    <div class="job-company">{job['company']} - {job['location']}</div>
                    <div style="font-size:12px; color:grey;">{t['source']}: {job['source']}</div>
                </div>
                """, unsafe_allow_html=True)
                st.link_button(t['apply'], job['link'])

if __name__ == "__main__":
    main()
